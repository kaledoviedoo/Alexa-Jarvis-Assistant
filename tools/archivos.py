"""
Herramientas de archivos: crear, leer, editar, mover, copiar, listar y buscar.

Todas las rutas pasan por `resolver_ruta`, que impide salir de las carpetas
permitidas en config.CARPETAS_PERMITIDAS. Nada de esto puede tocar C:\\Windows
aunque el modelo se confunda o alguien intente inyectar "../../".
"""

import logging
import os
import re
import shutil
from datetime import datetime
from pathlib import Path

from config import (
    CARPETAS_PERMITIDAS,
    DESCARGAS,
    DOCUMENTOS,
    ESCRITORIO,
    PAPELERA,
)

log = logging.getLogger("jarvis.archivos")


class RutaNoPermitida(Exception):
    """Se intentó tocar un archivo fuera de las carpetas autorizadas."""


# Alias hablados de carpetas -> ruta real
ALIAS_CARPETAS = {
    "escritorio": ESCRITORIO,
    "el escritorio": ESCRITORIO,
    "desktop": ESCRITORIO,
    "descargas": DESCARGAS,
    "las descargas": DESCARGAS,
    "downloads": DESCARGAS,
    "documentos": DOCUMENTOS,
    "los documentos": DOCUMENTOS,
    "documents": DOCUMENTOS,
}


def resolver_ruta(nombre: str, base: Path | None = None) -> Path:
    """
    Convierte un nombre hablado en una ruta absoluta segura.

    Acepta 'notas.txt', 'proyectos/notas.txt' o una ruta absoluta, siempre que
    el resultado quede dentro de una carpeta permitida.
    """
    nombre = (nombre or "").strip().strip('"').strip("'")
    if not nombre:
        raise RutaNoPermitida("No se indicó ningún nombre de archivo.")

    base = base or ESCRITORIO
    candidata = Path(nombre)
    es_absoluta = candidata.is_absolute()

    # Rechazo explícito de componentes de escalada ANTES de tocar el disco.
    #
    # No basta con confiar en resolve(): en Linux una cadena como '..\\..\\x'
    # no se interpreta como escalada (la barra invertida no es separador allí),
    # así que comprobamos ambos separadores a mano. Así el sandbox se comporta
    # igual en Windows, en WSL y en Linux.
    if not es_absoluta:
        # Separadores codificados en URL: nunca los decodificamos, así que no
        # escapan del sandbox, pero solo pueden venir de un intento de ataque.
        if re.search(r"%2[fF]|%5[cC]", nombre):
            raise RutaNoPermitida(f"'{nombre}' tiene caracteres codificados no permitidos.")

        # Letra de unidad (C:\...): en Windows pathlib ya la trata como
        # absoluta, pero en Linux o WSL quedaría como un nombre literal
        # absurdo dentro del Escritorio. La rechazamos en cualquier sistema.
        if re.match(r"^[a-zA-Z]:[\\/]", nombre):
            raise RutaNoPermitida(
                f"'{nombre}' apunta a una unidad del sistema y no está permitido."
            )

        componentes = re.split(r"[\\/]+", nombre)
        for componente in componentes:
            # '..' es escalada; '...' o '....' son nombres raros que no aportan
            # nada legítimo y que en algunos sistemas se normalizan de formas
            # inesperadas. Los bloqueamos todos.
            if componente and set(componente) == {"."} and len(componente) >= 2:
                raise RutaNoPermitida(
                    f"'{nombre}' contiene una ruta de escalada y no está permitida."
                )
        # Normalizamos las barras invertidas para que la unión con la base
        # funcione igual en cualquier sistema operativo.
        candidata = Path(nombre.replace("\\", "/"))

    ruta = candidata if es_absoluta else (base / candidata)

    # resolve() colapsa '..' y symlinks: es lo que hace segura la comprobación.
    try:
        ruta = ruta.resolve()
    except (OSError, RuntimeError) as e:
        raise RutaNoPermitida(f"Ruta inválida: {nombre}") from e

    for permitida in CARPETAS_PERMITIDAS:
        try:
            if ruta == permitida.resolve() or permitida.resolve() in ruta.parents:
                return ruta
        except OSError:
            continue

    raise RutaNoPermitida(
        f"'{nombre}' queda fuera de las carpetas autorizadas "
        f"(Escritorio, Descargas, Documentos)."
    )


def _base_desde_alias(carpeta: str | None) -> Path:
    if not carpeta:
        return ESCRITORIO
    clave = carpeta.strip().lower()
    return ALIAS_CARPETAS.get(clave, ESCRITORIO / carpeta)


# -------------------------------------------------------------------------
# CREAR
# -------------------------------------------------------------------------
def crear_archivo(nombre_archivo: str, contenido: str = "", carpeta: str = "") -> str:
    """Crea un archivo (.py, .txt, .md, .json, .csv, .docx, .xlsx) en el equipo."""
    try:
        ruta = resolver_ruta(nombre_archivo, _base_desde_alias(carpeta))
    except RutaNoPermitida as e:
        return str(e)

    ext = ruta.suffix.lower()
    if not ext:
        ruta = ruta.with_suffix(".txt")
        ext = ".txt"

    try:
        ruta.parent.mkdir(parents=True, exist_ok=True)

        if ext == ".docx":
            import docx  # import perezoso: solo si de verdad se necesita

            documento = docx.Document()
            for parrafo in (contenido or "").split("\n"):
                documento.add_paragraph(parrafo)
            documento.save(ruta)

        elif ext == ".xlsx":
            import openpyxl

            libro = openpyxl.Workbook()
            hoja = libro.active
            for i, linea in enumerate((contenido or "").split("\n"), start=1):
                # Las comas se reparten en columnas: útil para dictar tablas.
                for j, celda in enumerate(linea.split(","), start=1):
                    hoja.cell(row=i, column=j, value=celda.strip())
            libro.save(ruta)

        else:
            ruta.write_text(contenido or "", encoding="utf-8")

        log.info("Archivo creado: %s (%d caracteres)", ruta, len(contenido or ""))
        return f"Listo, creé {ruta.name} en {ruta.parent.name}."

    except ImportError as e:
        falta = "python-docx" if ext == ".docx" else "openpyxl"
        log.error("Falta dependencia para %s: %s", ext, e)
        return f"No pude crear el archivo: falta instalar {falta}."
    except Exception as e:
        log.exception("Error creando %s", ruta)
        return f"Hubo un error al crear {ruta.name}: {e}"


# -------------------------------------------------------------------------
# LEER
# -------------------------------------------------------------------------
def leer_archivo(nombre_archivo: str, carpeta: str = "") -> str:
    """Lee el contenido de un archivo de texto."""
    try:
        ruta = resolver_ruta(nombre_archivo, _base_desde_alias(carpeta))
    except RutaNoPermitida as e:
        return str(e)

    if not ruta.exists():
        return f"No encontré {ruta.name}."

    if ruta.suffix.lower() == ".docx":
        try:
            import docx

            documento = docx.Document(ruta)
            texto = "\n".join(p.text for p in documento.paragraphs)
        except Exception as e:
            return f"No pude leer el documento: {e}"
    else:
        try:
            texto = ruta.read_text(encoding="utf-8", errors="replace")
        except Exception as e:
            return f"No pude leer {ruta.name}: {e}"

    if not texto.strip():
        return f"{ruta.name} está vacío."

    return f"Contenido de {ruta.name}:\n{texto[:2000]}"


# -------------------------------------------------------------------------
# EDITAR
# -------------------------------------------------------------------------
def editar_archivo(
    nombre_archivo: str,
    accion: str = "agregar",
    contenido: str = "",
    buscar: str = "",
    reemplazar: str = "",
    carpeta: str = "",
) -> str:
    """
    Edita un archivo de texto existente.

    accion: 'agregar' (al final), 'anteponer' (al inicio),
            'reemplazar' (buscar->reemplazar) o 'sobrescribir'.
    """
    try:
        ruta = resolver_ruta(nombre_archivo, _base_desde_alias(carpeta))
    except RutaNoPermitida as e:
        return str(e)

    if not ruta.exists():
        return f"No encontré {ruta.name} para editar."

    if ruta.suffix.lower() in (".docx", ".xlsx"):
        return f"Todavía no puedo editar {ruta.suffix} por voz, solo archivos de texto."

    try:
        original = ruta.read_text(encoding="utf-8", errors="replace")
    except Exception as e:
        return f"No pude abrir {ruta.name}: {e}"

    accion = (accion or "agregar").strip().lower()

    if accion in ("agregar", "añadir", "anadir", "append"):
        separador = "" if original.endswith("\n") or not original else "\n"
        nuevo = original + separador + contenido
        resumen = f"Agregué el texto al final de {ruta.name}."

    elif accion in ("anteponer", "prepend", "inicio"):
        nuevo = contenido + "\n" + original
        resumen = f"Agregué el texto al inicio de {ruta.name}."

    elif accion in ("reemplazar", "replace", "sustituir"):
        if not buscar:
            return "Para reemplazar necesito saber qué texto buscar."
        if buscar not in original:
            return f"No encontré '{buscar}' dentro de {ruta.name}."
        veces = original.count(buscar)
        nuevo = original.replace(buscar, reemplazar)
        resumen = f"Reemplacé {veces} coincidencia{'s' if veces != 1 else ''} en {ruta.name}."

    elif accion in ("sobrescribir", "overwrite", "reescribir"):
        nuevo = contenido
        resumen = f"Reescribí {ruta.name} por completo."

    else:
        return f"No entendí la acción de edición '{accion}'."

    # Copia de seguridad antes de tocar nada.
    try:
        respaldo = PAPELERA / f"{ruta.stem}_{datetime.now():%Y%m%d_%H%M%S}{ruta.suffix}"
        shutil.copy2(ruta, respaldo)
        ruta.write_text(nuevo, encoding="utf-8")
    except Exception as e:
        log.exception("Error editando %s", ruta)
        return f"No pude guardar los cambios: {e}"

    log.info("Archivo editado: %s (%s)", ruta, accion)
    return resumen


# -------------------------------------------------------------------------
# MOVER / COPIAR / ELIMINAR
# -------------------------------------------------------------------------
def mover_archivo(origen: str, destino: str, carpeta: str = "") -> str:
    """Mueve un archivo a otra carpeta."""
    try:
        ruta_origen = resolver_ruta(origen, _base_desde_alias(carpeta))
    except RutaNoPermitida as e:
        return str(e)

    if not ruta_origen.exists():
        return f"No encontré {ruta_origen.name}."

    destino_limpio = (destino or "").strip().lower()
    if destino_limpio in ALIAS_CARPETAS:
        carpeta_destino = ALIAS_CARPETAS[destino_limpio]
    else:
        try:
            carpeta_destino = resolver_ruta(destino, ESCRITORIO)
        except RutaNoPermitida as e:
            return str(e)

    try:
        carpeta_destino.mkdir(parents=True, exist_ok=True)
        final = carpeta_destino / ruta_origen.name

        # No pisamos un archivo existente: le añadimos un sufijo.
        contador = 1
        while final.exists():
            final = carpeta_destino / f"{ruta_origen.stem}_{contador}{ruta_origen.suffix}"
            contador += 1

        shutil.move(str(ruta_origen), str(final))
    except Exception as e:
        log.exception("Error moviendo %s", ruta_origen)
        return f"No pude mover el archivo: {e}"

    log.info("Movido: %s -> %s", ruta_origen, final)
    return f"Moví {ruta_origen.name} a {carpeta_destino.name}."


def copiar_archivo(origen: str, destino: str, carpeta: str = "") -> str:
    """Copia un archivo a otra carpeta."""
    try:
        ruta_origen = resolver_ruta(origen, _base_desde_alias(carpeta))
    except RutaNoPermitida as e:
        return str(e)

    if not ruta_origen.exists():
        return f"No encontré {ruta_origen.name}."

    destino_limpio = (destino or "").strip().lower()
    carpeta_destino = ALIAS_CARPETAS.get(destino_limpio)
    if carpeta_destino is None:
        try:
            carpeta_destino = resolver_ruta(destino, ESCRITORIO)
        except RutaNoPermitida as e:
            return str(e)

    try:
        carpeta_destino.mkdir(parents=True, exist_ok=True)
        shutil.copy2(ruta_origen, carpeta_destino / ruta_origen.name)
    except Exception as e:
        return f"No pude copiar el archivo: {e}"

    return f"Copié {ruta_origen.name} a {carpeta_destino.name}."


def eliminar_archivo(nombre_archivo: str, carpeta: str = "") -> str:
    """
    Manda un archivo a la papelera interna de Jarvis.

    Nunca borra de forma definitiva: siempre se puede recuperar de ~/.jarvis/papelera.
    """
    try:
        ruta = resolver_ruta(nombre_archivo, _base_desde_alias(carpeta))
    except RutaNoPermitida as e:
        return str(e)

    if not ruta.exists():
        return f"No encontré {ruta.name}."

    try:
        marca = datetime.now().strftime("%Y%m%d_%H%M%S")
        shutil.move(str(ruta), str(PAPELERA / f"{marca}_{ruta.name}"))
    except Exception as e:
        return f"No pude eliminar el archivo: {e}"

    log.info("Enviado a papelera: %s", ruta)
    return f"Mandé {ruta.name} a la papelera de Jarvis. Se puede recuperar."


# -------------------------------------------------------------------------
# LISTAR / BUSCAR
# -------------------------------------------------------------------------
def listar_archivos(carpeta: str = "escritorio", limite: int = 15) -> str:
    """Lista los archivos de una carpeta."""
    base = _base_desde_alias(carpeta)
    try:
        base = resolver_ruta(str(base), base.parent if base.parent.exists() else base)
    except RutaNoPermitida as e:
        return str(e)

    if not base.is_dir():
        return f"No encontré la carpeta {carpeta}."

    try:
        entradas = sorted(
            (p for p in base.iterdir() if not p.name.startswith(".")),
            key=lambda p: p.stat().st_mtime,
            reverse=True,
        )
    except Exception as e:
        return f"No pude leer la carpeta: {e}"

    if not entradas:
        return f"La carpeta {base.name} está vacía."

    nombres = [p.name for p in entradas[:limite]]
    total = len(entradas)
    resumen = ", ".join(nombres)

    if total > limite:
        return f"Hay {total} elementos en {base.name}. Los más recientes: {resumen}."
    return f"En {base.name} hay {total}: {resumen}."


def buscar_archivo(patron: str, carpeta: str = "", limite: int = 10) -> str:
    """Busca archivos por nombre dentro de las carpetas permitidas."""
    patron = (patron or "").strip().lower()
    if not patron:
        return "¿Qué archivo quieres que busque?"

    bases = [_base_desde_alias(carpeta)] if carpeta else CARPETAS_PERMITIDAS
    encontrados: list[Path] = []

    for base in bases:
        if not base.is_dir():
            continue
        try:
            for raiz, dirs, ficheros in os.walk(base):
                # No entramos en carpetas ocultas ni pesadas.
                dirs[:] = [
                    d for d in dirs
                    if not d.startswith(".") and d not in ("node_modules", "__pycache__", "venv")
                ]
                for fichero in ficheros:
                    if patron in fichero.lower():
                        encontrados.append(Path(raiz) / fichero)
                        if len(encontrados) >= limite:
                            break
                if len(encontrados) >= limite:
                    break
        except Exception:
            continue
        if len(encontrados) >= limite:
            break

    if not encontrados:
        return f"No encontré ningún archivo que contenga '{patron}'."

    if len(encontrados) == 1:
        p = encontrados[0]
        return f"Encontré {p.name} en {p.parent.name}."

    nombres = ", ".join(p.name for p in encontrados[:5])
    return f"Encontré {len(encontrados)} archivos: {nombres}."


def crear_carpeta(nombre: str, carpeta: str = "") -> str:
    """Crea una carpeta nueva."""
    try:
        ruta = resolver_ruta(nombre, _base_desde_alias(carpeta))
    except RutaNoPermitida as e:
        return str(e)

    try:
        ruta.mkdir(parents=True, exist_ok=True)
    except Exception as e:
        return f"No pude crear la carpeta: {e}"

    return f"Creé la carpeta {ruta.name}."


def eliminar_varios(patron: str, carpeta: str = "", limite: int = 20) -> str:
    """
    Manda a la papelera todos los archivos cuyo nombre contenga el patron.

    Pensado para ordenes como "elimina las capturas del escritorio". El limite
    existe a proposito: si una orden mal entendida fuera a barrer media carpeta,
    preferimos no hacerlo y decirlo.
    """
    patron = (patron or "").strip().lower()
    if not patron:
        return "¿Qué archivos quieres que elimine?"

    base = _base_desde_alias(carpeta)
    if not base.is_dir():
        return f"No encontré la carpeta {carpeta or 'escritorio'}."

    try:
        candidatos = [
            p for p in base.iterdir()
            if p.is_file() and patron in p.name.lower()
        ]
    except Exception as e:
        return f"No pude leer la carpeta: {e}"

    if not candidatos:
        return f"No encontré ningún archivo que contenga '{patron}'."

    if len(candidatos) > limite:
        return (
            f"Encontré {len(candidatos)} archivos con '{patron}'. "
            f"Son demasiados para borrarlos de una vez: hazlo desde el explorador."
        )

    movidos = []
    for ruta in candidatos:
        try:
            marca = datetime.now().strftime("%Y%m%d_%H%M%S")
            shutil.move(str(ruta), str(PAPELERA / f"{marca}_{ruta.name}"))
            movidos.append(ruta.name)
        except Exception as e:
            log.warning("No pude mover %s: %s", ruta, e)

    if not movidos:
        return "No pude eliminar ninguno de esos archivos."

    log.info("Eliminados %d archivos con patron %r", len(movidos), patron)

    if len(movidos) == 1:
        return f"Mandé {movidos[0]} a la papelera de Jarvis."
    return f"Mandé {len(movidos)} archivos a la papelera de Jarvis."
